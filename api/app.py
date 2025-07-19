from flask import Flask, render_template, request, send_file
from docx import Document
import os
from datetime import datetime
import zipfile
import logging
import re
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

app = Flask(__name__, static_folder="../static", template_folder="../templates")

# Define directories using /tmp/ for Vercel compatibility
LOGS_DIR = "/tmp/logs"
OUTPUT_DIR = "/tmp/output"
DOWNLOAD_DIR = "/tmp/downloads"
TEMPLATES_DIR = "templates_docx"

# Create directories if they don't exist
for directory in [LOGS_DIR, OUTPUT_DIR, DOWNLOAD_DIR]:
    os.makedirs(directory, exist_ok=True)

# Setup logging
logging.basicConfig(
    filename=os.path.join(LOGS_DIR, 'app.log'),
    level=logging.INFO,
    format='%(asctime)s %(levelname)s: %(message)s'
)

# Setup rate limiting
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://"
)

@app.route('/')
@limiter.limit("10 per minute")
def home():
    return render_template('home.html')

@app.route('/sperm')
@limiter.limit("10 per minute")
def sperm_index():
    form_data = {
        'full_name': '', 'address': '', 'pin_code': '', 'contact_number': '', 'aadhaar_number': '',
        'date_of_birth': '', 'email_address': '', 'donor_id': '', 'date_of_discussion': '',
        'date_of_consultancy': '', 'genetic_disorders': '', 'family_history': '', 'current_medications': '',
        'allergies': '', 'last_medical_exam': '', 'hiv_results': '', 'hbv_results': '', 'hcv_results': '',
        'vdrl_results': '', 'serious_illness': '', 'smoking': 'No', 'smoking_frequency': '',
        'cigarettes_per_day': '', 'alcohol': 'No', 'alcohol_frequency': '', 'alcohol_amount': '',
        'drug_use': 'No', 'diet': 'Not Specified', 'marital_status': 'Not Specified', 'num_children': '0',
        'donor_experience': 'No', 'donation_frequency': '', 'height': '', 'weight': '', 'education': '',
        'mother_tongue': '', 'skin_colour': '', 'hair_colour': '', 'eye_colour': '', 'religion': '',
        'occupation': '', 'consent_cryopreservation': 'No', 'consent_art_bank': 'No', 'consent_registry': 'No'
    }
    return render_template('sperm_index.html', form_data=form_data, errors=[], today=datetime.now().date().isoformat())

@app.route('/oocyte')
@limiter.limit("10 per minute")
def oocyte_index():
    form_data = {
        'full_name': '', 'address': '', 'district': '', 'state': '', 'pin_code': '', 'contact_number': '',
        'aadhaar_number': '', 'date_of_birth': '', 'age': '', 'date_of_discussion': '', 'date_of_consultancy': '',
        'marital_status': 'Not Specified', 'num_children': '0', 'donor_id': '', 'last_medical_exam': '',
        'hiv_results': '', 'hbv_results': '', 'hcv_results': '', 'vdrl_results': '', 'family_history': '',
        'serious_illness': '', 'current_medications': '', 'allergies': '', 'antral_follicle_count': '',
        'fsh_levels': '', 'amh_levels': '', 'tobacco_use': 'No', 'tobacco_frequency': '', 'alcohol': 'No',
        'alcohol_frequency': '', 'drug_use': 'No', 'exercise_routine': '', 'consent_registry': 'No',
        'place': '', 'ivf_name': '', 'ivf_address': '', 'doctor_name': ''
    }
    return render_template('oocyte_index.html', form_data=form_data, errors=[], today=datetime.now().date().isoformat())

@app.route('/commissioning_couple')
@limiter.limit("10 per minute")
def commissioning_couple_index():
    form_data = {
        'female_name': '', 'male_name': '', 'female_age': '', 'female_dob': '', 'female_aadhaar': '',
        'female_occupation': '', 'male_age': '', 'male_dob': '', 'address': '', 'district': '',
        'state': '', 'pin_code': '', 'place': '', 'ivf_name': '', 'ivf_address': '', 'doctor_name': ''
    }
    return render_template('commissioning_couple_index.html', form_data=form_data, errors=[], today=datetime.now().date().isoformat())

@app.route('/generate_sperm', methods=['POST'])
@limiter.limit("5 per minute")
def generate_sperm_document():
    logging.info(f"Processing sperm donor form submission from {request.remote_addr}")
    form_data = {
        'full_name': request.form.get('full_name', '').strip(),
        'address': request.form.get('address', '').strip(),
        'pin_code': request.form.get('pin_code', '').strip(),
        'contact_number': request.form.get('contact_number', '').strip(),
        'aadhaar_number': request.form.get('aadhaar_number', '').strip(),
        'date_of_birth': request.form.get('date_of_birth', '').strip(),
        'email_address': request.form.get('email_address', '').strip(),
        'donor_id': request.form.get('donor_id', '').strip(),
        'date_of_discussion': request.form.get('date_of_discussion', '').strip(),
        'date_of_consultancy': request.form.get('date_of_consultancy', '').strip(),
        'genetic_disorders': request.form.get('genetic_disorders', '').strip(),
        'family_history': request.form.get('family_history', '').strip(),
        'current_medications': request.form.get('current_medications', '').strip(),
        'allergies': request.form.get('allergies', '').strip(),
        'last_medical_exam': request.form.get('last_medical_exam', '').strip(),
        'hiv_results': request.form.get('hiv_results', '').strip(),
        'hbv_results': request.form.get('hbv_results', '').strip(),
        'hcv_results': request.form.get('hcv_results', '').strip(),
        'vdrl_results': request.form.get('vdrl_results', '').strip(),
        'serious_illness': request.form.get('serious_illness', '').strip(),
        'smoking': request.form.get('smoking', 'No'),
        'smoking_frequency': request.form.get('smoking_frequency', '').strip(),
        'cigarettes_per_day': request.form.get('cigarettes_per_day', '').strip(),
        'alcohol': request.form.get('alcohol', 'No'),
        'alcohol_frequency': request.form.get('alcohol_frequency', '').strip(),
        'alcohol_amount': request.form.get('alcohol_amount', '').strip(),
        'drug_use': request.form.get('drug_use', 'No'),
        'diet': request.form.get('diet', 'Not Specified'),
        'marital_status': request.form.get('marital_status', 'Not Specified'),
        'num_children': request.form.get('num_children', '0').strip(),
        'donor_experience': request.form.get('donor_experience', 'No'),
        'donation_frequency': request.form.get('donation_frequency', '').strip(),
        'height': request.form.get('height', '').strip(),
        'weight': request.form.get('weight', '').strip(),
        'education': request.form.get('education', '').strip(),
        'mother_tongue': request.form.get('mother_tongue', '').strip(),
        'skin_colour': request.form.get('skin_colour', '').strip(),
        'hair_colour': request.form.get('hair_colour', '').strip(),
        'eye_colour': request.form.get('eye_colour', '').strip(),
        'religion': request.form.get('religion', '').strip(),
        'occupation': request.form.get('occupation', '').strip(),
        'consent_cryopreservation': 'Yes' if request.form.get('consent_cryopreservation') else 'No',
        'consent_art_bank': 'Yes' if request.form.get('consent_art_bank') else 'No',
        'consent_registry': 'Yes' if request.form.get('consent_registry') else 'No',
        'date': datetime.now().strftime('%d/%m/%y')
    }

    # Server-side validation
    required_fields = ['full_name', 'aadhaar_number', 'date_of_discussion', 'date_of_consultancy']
    errors = []
    for field in required_fields:
        if not form_data[field]:
            errors.append(f"{field.replace('_', ' ').title()} is required.")

    # Validate Aadhaar number
    if form_data['aadhaar_number'] and not re.match(r'^\d{12}$', form_data['aadhaar_number']):
        errors.append("Aadhaar Number must be 12 digits.")

    # Validate email format
    if form_data['email_address'] and not re.match(r'^[\w\.-]+@[\w\.-]+\.\w+$', form_data['email_address']):
        errors.append("Invalid email address.")

    # Validate phone number
    if form_data['contact_number'] and not re.match(r'^\d{10}$|^$', form_data['contact_number']):
        errors.append("Contact Number must be 10 digits if provided.")

    # Validate PIN code
    if form_data['pin_code'] and not re.match(r'^\d{6}$|^$', form_data['pin_code']):
        errors.append("PIN Code must be 6 digits if provided.")

    # Validate dates
    today = datetime.now().date()
    for date_field in ['date_of_birth', 'last_medical_exam', 'date_of_discussion', 'date_of_consultancy']:
        if form_data[date_field]:
            try:
                input_date = datetime.strptime(form_data[date_field], '%Y-%m-%d').date()
                if input_date > today:
                    errors.append(f"{date_field.replace('_', ' ').title()} cannot be in the future.")
            except ValueError:
                errors.append(f"Invalid format for {date_field.replace('_', ' ').title()}.")

    # Validate number of children
    try:
        num_children = int(form_data['num_children']) if form_data['num_children'] else 0
        if num_children < 0:
            errors.append("Number of children cannot be negative.")
        elif num_children > 20:
            errors.append("Number of children cannot exceed 20.")
    except ValueError:
        errors.append("Number of children must be a valid number.")

    # Validate children ages
    children_ages = []
    for i in range(1, num_children + 1):
        age = request.form.get(f'child_{i}_age', '').strip()
        try:
            if age:
                age_val = int(age)
                if age_val < 0 or age_val > 100:
                    errors.append(f"Child {i} Age must be between 0 and 100.")
                children_ages.append(f"Child {i}: Age: {age}")
            else:
                children_ages.append(f"Child {i}: Age: N/A")
        except ValueError:
            errors.append(f"Child {i} Age must be a valid number.")

    # Validate conditional fields
    if form_data['smoking'] == 'Yes' and not (form_data['smoking_frequency'] and form_data['cigarettes_per_day']):
        errors.append("Smoking Frequency and Cigarettes per Day are required if smoking is Yes.")
    if form_data['alcohol'] == 'Yes' and not (form_data['alcohol_frequency'] and form_data['alcohol_amount']):
        errors.append("Alcohol Frequency and Amount are required if alcohol consumption is Yes.")

    # Validate blood test results
    for field in ['hiv_results', 'hbv_results', 'hcv_results', 'vdrl_results']:
        if form_data[field] and form_data[field].lower() not in ['negative', 'positive', 'pending', '']:
            errors.append(f"{field.replace('_', ' ').title()} must be 'Negative', 'Positive', or 'Pending'.")

    if errors:
        logging.warning(f"Validation errors: {errors}")
        return render_template('sperm_index.html', errors=errors, form_data=form_data, today=datetime.now().date().isoformat(), show_modal=True)

    form_data['children_ages'] = '\n'.join(children_ages) if children_ages else 'None'

    # List of templates and their required fields
    templates = [
        {
            'file': os.path.join(TEMPLATES_DIR, 'form_15.docx'),
            'output': 'form_15_{full_name}_{timestamp}.docx',
            'fields': ['full_name', 'address', 'pin_code', 'contact_number', 'aadhaar_number', 'date_of_discussion', 'date_of_consultancy', 'date']
        },
        {
            'file': os.path.join(TEMPLATES_DIR, 'medical_history.docx'),
            'output': 'medical_history_{full_name}_{timestamp}.docx',
            'fields': ['full_name', 'date_of_birth', 'address', 'contact_number', 'email_address', 'aadhaar_number', 'donor_id', 'date', 'last_medical_exam', 'hiv_results', 'hbv_results', 'hcv_results', 'vdrl_results', 'family_history', 'serious_illness', 'current_medications', 'allergies', 'consent_cryopreservation', 'consent_art_bank', 'consent_registry']
        },
        {
            'file': os.path.join(TEMPLATES_DIR, 'donor_info.docx'),
            'output': 'donor_info_{full_name}_{timestamp}.docx',
            'fields': ['full_name', 'date_of_birth', 'contact_number', 'email_address', 'aadhaar_number', 'genetic_disorders', 'family_history', 'current_medications', 'allergies', 'smoking', 'smoking_frequency', 'cigarettes_per_day', 'alcohol', 'alcohol_frequency', 'alcohol_amount', 'drug_use', 'diet', 'marital_status', 'num_children', 'donor_experience', 'donation_frequency', 'height', 'weight', 'education', 'mother_tongue', 'skin_colour', 'hair_colour', 'eye_colour', 'religion', 'occupation', 'date', 'children_ages']
        }
    ]

    # Generate documents and prepare ZIP
    generated_files = []
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    zip_filename = f"sperm_donor_documents_{form_data['full_name'].replace(' ', '_') or 'Unknown'}_{timestamp}.zip"
    zip_path = os.path.join(DOWNLOAD_DIR, zip_filename)

    try:
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for template in templates:
                try:
                    doc = Document(template['file'])
                    for paragraph in doc.paragraphs:
                        for key in template['fields']:
                            placeholder = f'{{{key}}}'
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, form_data[key] or 'N/A')
                    output_filename = template['output'].format(full_name=form_data['full_name'].replace(' ', '_') or 'Unknown', timestamp=timestamp)
                    output_path = os.path.join(OUTPUT_DIR, output_filename)
                    doc.save(output_path)
                    zipf.write(output_path, output_filename)
                    generated_files.append(output_path)
                    logging.info(f"Generated {output_filename}")
                except FileNotFoundError as e:
                    logging.error(f"Template file {template['file']} not found: {str(e)}")
                    return render_template('error.html', errors=[f"Template file {template['file']} not found."], show_modal=True)
                except Exception as e:
                    logging.error(f"Error processing {template['file']}: {str(e)}")
                    return render_template('error.html', errors=[f"Error processing document: {str(e)}"], show_modal=True)
    except Exception as e:
        logging.error(f"Error creating ZIP file: {str(e)}")
        return render_template('error.html', errors=[f"Error creating ZIP file: {str(e)}"], show_modal=True)

    logging.info(f"ZIP file created: {zip_filename}")
    return render_template('success.html', zip_filename=zip_filename, errors=[], donor_type='sperm')

@app.route('/generate_oocyte', methods=['POST'])
@limiter.limit("5 per minute")
def generate_oocyte_document():
    logging.info(f"Processing oocyte donor form submission from {request.remote_addr}")
    form_data = {
        'full_name': request.form.get('full_name', '').strip(),
        'address': request.form.get('address', '').strip(),
        'district': request.form.get('district', '').strip(),
        'state': request.form.get('state', '').strip(),
        'pin_code': request.form.get('pin_code', '').strip(),
        'contact_number': request.form.get('contact_number', '').strip(),
        'aadhaar_number': request.form.get('aadhaar_number', '').strip(),
        'date_of_birth': request.form.get('date_of_birth', '').strip(),
        'age': request.form.get('age', '').strip(),
        'date_of_discussion': request.form.get('date_of_discussion', '').strip(),
        'date_of_consultancy': request.form.get('date_of_consultancy', '').strip(),
        'marital_status': request.form.get('marital_status', 'Not Specified'),
        'num_children': request.form.get('num_children', '0').strip(),
        'donor_id': request.form.get('donor_id', '').strip(),
        'last_medical_exam': request.form.get('last_medical_exam', '').strip(),
        'hiv_results': request.form.get('hiv_results', '').strip(),
        'hbv_results': request.form.get('hbv_results', '').strip(),
        'hcv_results': request.form.get('hcv_results', '').strip(),
        'vdrl_results': request.form.get('vdrl_results', '').strip(),
        'family_history': request.form.get('family_history', '').strip(),
        'serious_illness': request.form.get('serious_illness', '').strip(),
        'current_medications': request.form.get('current_medications', '').strip(),
        'allergies': request.form.get('allergies', '').strip(),
        'antral_follicle_count': request.form.get('antral_follicle_count', '').strip(),
        'fsh_levels': request.form.get('fsh_levels', '').strip(),
        'amh_levels': request.form.get('amh_levels', '').strip(),
        'tobacco_use': request.form.get('tobacco_use', 'No'),
        'tobacco_frequency': request.form.get('tobacco_frequency', '').strip(),
        'alcohol': request.form.get('alcohol', 'No'),
        'alcohol_frequency': request.form.get('alcohol_frequency', '').strip(),
        'drug_use': request.form.get('drug_use', 'No'),
        'exercise_routine': request.form.get('exercise_routine', '').strip(),
        'consent_registry': 'Yes' if request.form.get('consent_registry') else 'No',
        'place': request.form.get('place', '').strip(),
        'ivf_name': request.form.get('ivf_name', '').strip(),
        'ivf_address': request.form.get('ivf_address', '').strip(),
        'doctor_name': request.form.get('doctor_name', '').strip(),
        'date': datetime.now().strftime('%d/%m/%y')
    }

    # Server-side validation
    required_fields = ['full_name', 'aadhaar_number', 'date_of_discussion', 'date_of_consultancy', 'ivf_name', 'ivf_address', 'doctor_name']
    errors = []
    for field in required_fields:
        if not form_data[field]:
            errors.append(f"{field.replace('_', ' ').title()} is required.")

    # Validate Aadhaar number
    if form_data['aadhaar_number'] and not re.match(r'^\d{12}$', form_data['aadhaar_number']):
        errors.append("Aadhaar Number must be 12 digits.")

    # Validate phone number
    if form_data['contact_number'] and not re.match(r'^\d{10}$|^$', form_data['contact_number']):
        errors.append("Contact Number must be 10 digits if provided.")

    # Validate PIN code
    if form_data['pin_code'] and not re.match(r'^\d{6}$|^$', form_data['pin_code']):
        errors.append("PIN Code must be 6 digits if provided.")

    # Validate age
    try:
        age = int(form_data['age']) if form_data['age'] else 0
        if age < 18 or age > 100:
            errors.append("Age must be between 18 and 100.")
    except ValueError:
        errors.append("Age must be a valid number.")

    # Validate dates
    today = datetime.now().date()
    for date_field in ['date_of_birth', 'date_of_discussion', 'date_of_consultancy', 'last_medical_exam']:
        if form_data[date_field]:
            try:
                input_date = datetime.strptime(form_data[date_field], '%Y-%m-%d').date()
                if input_date > today:
                    errors.append(f"{date_field.replace('_', ' ').title()} cannot be in the future.")
            except ValueError:
                errors.append(f"Invalid format for {date_field.replace('_', ' ').title()}.")

    # Validate number of children
    try:
        num_children = int(form_data['num_children']) if form_data['num_children'] else 0
        if num_children < 0:
            errors.append("Number of children cannot be negative.")
        elif num_children > 20:
            errors.append("Number of children cannot exceed 20.")
    except ValueError:
        errors.append("Number of children must be a valid number.")

    # Validate children ages
    children_ages = []
    for i in range(1, num_children + 1):
        age = request.form.get(f'child_{i}_age', '').strip()
        try:
            if age:
                age_val = int(age)
                if age_val < 2 or age_val > 100:
                    errors.append(f"Child {i} Age must be between 2 and 100.")
                children_ages.append(f"Child {i}: Age: {age}")
            else:
                children_ages.append(f"Child {i}: Age: N/A")
        except ValueError:
            errors.append(f"Child {i} Age must be a valid number.")

    # Validate conditional fields
    if form_data['tobacco_use'] == 'Yes' and not form_data['tobacco_frequency']:
        errors.append("Tobacco Frequency is required if tobacco use is Yes.")
    if form_data['alcohol'] == 'Yes' and not form_data['alcohol_frequency']:
        errors.append("Alcohol Frequency is required if alcohol consumption is Yes.")

    # Validate blood test results
    for field in ['hiv_results', 'hbv_results', 'hcv_results', 'vdrl_results']:
        if form_data[field] and form_data[field].lower() not in ['negative', 'positive', 'pending', '']:
            errors.append(f"{field.replace('_', ' ').title()} must be 'Negative', 'Positive', or 'Pending'.")

    if errors:
        logging.warning(f"Validation errors: {errors}")
        return render_template('oocyte_index.html', errors=errors, form_data=form_data, today=datetime.now().date().isoformat(), show_modal=True)

    form_data['children_ages'] = '\n'.join(children_ages) if children_ages else 'None'

    # List of templates and their required fields
    templates = [
        {
            'file': os.path.join(TEMPLATES_DIR, 'form_13.docx'),
            'output': 'form_13_{full_name}_{timestamp}.docx',
            'fields': ['full_name', 'address', 'district', 'state', 'pin_code', 'contact_number', 'aadhaar_number', 'date_of_discussion', 'date', 'ivf_name', 'ivf_address', 'doctor_name']
        },
        {
            'file': os.path.join(TEMPLATES_DIR, 'oocyte_medical_history.docx'),
            'output': 'oocyte_medical_history_{full_name}_{timestamp}.docx',
            'fields': ['full_name', 'date_of_birth', 'marital_status', 'address', 'district', 'state', 'pin_code', 'contact_number', 'email_address', 'aadhaar_number', 'donor_id', 'date', 'last_medical_exam', 'hiv_results', 'hbv_results', 'hcv_results', 'vdrl_results', 'family_history', 'serious_illness', 'current_medications', 'allergies', 'antral_follicle_count', 'fsh_levels', 'amh_levels', 'tobacco_use', 'tobacco_frequency', 'alcohol', 'alcohol_frequency', 'drug_use', 'exercise_routine', 'consent_registry']
        },
        {
            'file': os.path.join(TEMPLATES_DIR, 'oocyte_donor_affidavit.docx'),
            'output': 'oocyte_donor_affidavit_{full_name}_{timestamp}.docx',
            'fields': ['full_name', 'age', 'date_of_birth', 'address', 'district', 'state', 'pin_code', 'contact_number', 'aadhaar_number', 'num_children', 'place', 'date', 'ivf_name', 'ivf_address', 'doctor_name']
        }
    ]

    # Generate documents and prepare ZIP
    generated_files = []
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    zip_filename = f"oocyte_donor_documents_{form_data['full_name'].replace(' ', '_') or 'Unknown'}_{timestamp}.zip"
    zip_path = os.path.join(DOWNLOAD_DIR, zip_filename)

    try:
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for template in templates:
                try:
                    doc = Document(template['file'])
                    for paragraph in doc.paragraphs:
                        for key in template['fields']:
                            placeholder = f'{{{key}}}'
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, form_data[key] or 'N/A')
                    output_filename = template['output'].format(full_name=form_data['full_name'].replace(' ', '_') or 'Unknown', timestamp=timestamp)
                    output_path = os.path.join(OUTPUT_DIR, output_filename)
                    doc.save(output_path)
                    zipf.write(output_path, output_filename)
                    generated_files.append(output_path)
                    logging.info(f"Generated {output_filename}")
                except FileNotFoundError as e:
                    logging.error(f"Template file {template['file']} not found: {str(e)}")
                    return render_template('error.html', errors=[f"Template file {template['file']} not found."], show_modal=True)
                except Exception as e:
                    logging.error(f"Error processing {template['file']}: {str(e)}")
                    return render_template('error.html', errors=[f"Error processing document: {str(e)}"], show_modal=True)
    except Exception as e:
        logging.error(f"Error creating ZIP file: {str(e)}")
        return render_template('error.html', errors=[f"Error creating ZIP file: {str(e)}"], show_modal=True)

    logging.info(f"ZIP file created: {zip_filename}")
    return render_template('success.html', zip_filename=zip_filename, errors=[], donor_type='oocyte')

@app.route('/generate_commissioning_couple', methods=['POST'])
@limiter.limit("5 per minute")
def generate_commissioning_couple_document():
    logging.info(f"Processing commissioning couple form submission from {request.remote_addr}")
    form_data = {
        'female_name': request.form.get('female_name', '').strip(),
        'male_name': request.form.get('male_name', '').strip(),
        'female_age': request.form.get('female_age', '').strip(),
        'female_dob': request.form.get('female_dob', '').strip(),
        'female_aadhaar': request.form.get('female_aadhaar', '').strip(),
        'female_occupation': request.form.get('female_occupation', '').strip(),
        'male_age': request.form.get('male_age', '').strip(),
        'male_dob': request.form.get('male_dob', '').strip(),
        'address': request.form.get('address', '').strip(),
        'district': request.form.get('district', '').strip(),
        'state': request.form.get('state', '').strip(),
        'pin_code': request.form.get('pin_code', '').strip(),
        'place': request.form.get('place', '').strip(),
        'ivf_name': request.form.get('ivf_name', '').strip(),
        'ivf_address': request.form.get('ivf_address', '').strip(),
        'doctor_name': request.form.get('doctor_name', '').strip(),
        'date': datetime.now().strftime('%d/%m/%y')
    }

    # Server-side validation
    required_fields = ['female_name', 'male_name', 'female_aadhaar', 'female_dob', 'male_dob', 'ivf_name', 'ivf_address', 'doctor_name']
    errors = []
    for field in required_fields:
        if not form_data[field]:
            errors.append(f"{field.replace('_', ' ').title()} is required.")

    # Validate Aadhaar number
    if form_data['female_aadhaar'] and not re.match(r'^\d{12}$', form_data['female_aadhaar']):
        errors.append("Female Aadhaar Number must be 12 digits.")

    # Validate PIN code
    if form_data['pin_code'] and not re.match(r'^\d{6}$|^$', form_data['pin_code']):
        errors.append("PIN Code must be 6 digits if provided.")

    # Validate ages
    for field in ['female_age', 'male_age']:
        try:
            age = int(form_data[field]) if form_data[field] else 0
            if age < 18 or age > 100:
                errors.append(f"{field.replace('_', ' ').title()} must be between 18 and 100.")
        except ValueError:
            errors.append(f"{field.replace('_', ' ').title()} must be a valid number.")

    # Validate dates
    today = datetime.now().date()
    for date_field in ['female_dob', 'male_dob']:
        if form_data[date_field]:
            try:
                input_date = datetime.strptime(form_data[date_field], '%Y-%m-%d').date()
                if input_date > today:
                    errors.append(f"{date_field.replace('_', ' ').title()} cannot be in the future.")
            except ValueError:
                errors.append(f"Invalid format for {date_field.replace('_', ' ').title()}.")

    if errors:
        logging.warning(f"Validation errors: {errors}")
        return render_template('commissioning_couple_index.html', errors=errors, form_data=form_data, today=datetime.now().date().isoformat(), show_modal=True)

    # List of templates and their required fields
    templates = [
        {
            'file': os.path.join(TEMPLATES_DIR, 'commissioning_couple_affidavit.docx'),
            'output': 'commissioning_couple_affidavit_{female_name}_{timestamp}.docx',
            'fields': ['female_name', 'male_name', 'female_age', 'female_dob', 'female_aadhaar', 'female_occupation', 'male_age', 'male_dob', 'address', 'district', 'state', 'pin_code', 'place', 'date', 'ivf_name', 'ivf_address', 'doctor_name']
        }
    ]

    # Generate documents and prepare ZIP
    generated_files = []
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    zip_filename = f"commissioning_couple_documents_{form_data['female_name'].replace(' ', '_') or 'Unknown'}_{timestamp}.zip"
    zip_path = os.path.join(DOWNLOAD_DIR, zip_filename)

    try:
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for template in templates:
                try:
                    doc = Document(template['file'])
                    for paragraph in doc.paragraphs:
                        for key in template['fields']:
                            placeholder = f'{{{key}}}'
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, form_data[key] or 'N/A')
                    output_filename = template['output'].format(female_name=form_data['female_name'].replace(' ', '_') or 'Unknown', timestamp=timestamp)
                    output_path = os.path.join(OUTPUT_DIR, output_filename)
                    doc.save(output_path)
                    zipf.write(output_path, output_filename)
                    generated_files.append(output_path)
                    logging.info(f"Generated {output_filename}")
                except FileNotFoundError as e:
                    logging.error(f"Template file {template['file']} not found: {str(e)}")
                    return render_template('error.html', errors=[f"Template file {template['file']} not found."], show_modal=True)
                except Exception as e:
                    logging.error(f"Error processing {template['file']}: {str(e)}")
                    return render_template('error.html', errors=[f"Error processing document: {str(e)}"], show_modal=True)
    except Exception as e:
        logging.error(f"Error creating ZIP file: {str(e)}")
        return render_template('error.html', errors=[f"Error creating ZIP file: {str(e)}"], show_modal=True)

    logging.info(f"ZIP file created: {zip_filename}")
    return render_template('success.html', zip_filename=zip_filename, errors=[], donor_type='commissioning_couple')

@app.route('/download/<filename>')
@limiter.limit("5 per minute")
def download_zip(filename):
    zip_path = os.path.join(DOWNLOAD_DIR, filename)
    try:
        response = send_file(zip_path, as_attachment=True)
        logging.info(f"Downloaded {filename} by {request.remote_addr}")
        # Clean up output files
        for file in os.listdir(OUTPUT_DIR):
            try:
                os.remove(os.path.join(OUTPUT_DIR, file))
            except Exception as e:
                logging.warning(f"Failed to delete output file {file}: {str(e)}")
        # Clean up ZIP file
        try:
            os.remove(zip_path)
        except Exception as e:
            logging.warning(f"Failed to delete ZIP file {zip_path}: {str(e)}")
        return response
    except FileNotFoundError:
        logging.error(f"ZIP file {filename} not found")
        return render_template('error.html', errors=["ZIP file not found. It may have already been downloaded or deleted."], show_modal=True)
    except Exception as e:
        logging.error(f"Error downloading {filename}: {str(e)}")
        return render_template('error.html', errors=[f"Error downloading file: {str(e)}"], show_modal=True)

@app.errorhandler(429)
def ratelimit_handler(e):
    logging.warning(f"Rate limit exceeded for {request.remote_addr}: {str(e)}")
    return render_template('error.html', errors=["Too many requests. Please try again later."], show_modal=True), 429

if __name__ == '__main__':
    app.run(debug=True)