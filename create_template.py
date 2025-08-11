from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def apply_official_style(doc):
    """
    Applies a consistent, official-looking style to the document.
    - Font: Times New Roman, 12pt
    - Body Alignment: Justified
    - Paragraph Spacing: 6pt after
    """
    # Set the default style for all paragraphs in the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.space_after = Pt(6)

    # Apply the same font to all heading styles
    for i in range(1, 3):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)

def create_form_15_template():
    """Generates the Sperm Donor Consent Form (Form 15) with official styling."""
    doc = Document()
    apply_official_style(doc)

    # --- Document Content ---
    heading = doc.add_heading('FORM 15', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].bold = True

    doc.add_paragraph('[Refer Rule 13 (2) (ii)]').alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_heading = doc.add_heading('Consent Form for the Donor of Sperm', level=2)
    sub_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('I, ').bold = True
    p.add_run('{full_name}, residing at {address}, PIN Code: {pin_code}, Mobile: {contact_number}, Aadhaar Number {aadhaar_number}, willingly consent to donate my sperm to couple/individual who are unable to have a child by other means. At this stage and to the best of my knowledge I am free of any infectious diseases or genetic disorders.')

    doc.add_paragraph('I have had a full discussion with ').add_run('Dr. Ravikumar N.R').bold = True
    doc.add_paragraph('on {date_of_discussion}, address Subash Nagara, B.H Road, Nelamangala Town, Bengaluru District, Karnataka-562123.')
    
    p = doc.add_paragraph('I have been counselled by ')
    p.add_run('Ranjana Basavaraj Byadagi').bold = True
    p.add_run(', address Avaraguppa, Avaraguppa Post, Siddapura, Uttara Kannada-581355 on {date_of_consultancy}.')
    
    doc.add_paragraph('(I understand that there will be no direct or indirect contact between the recipient, and me, and my personal identity will not be disclosed to the recipient or to the child born through the use of my gamete: If applicable)')
    doc.add_paragraph('I understand that I shall have no rights whatsoever on the resulting offspring and vice versa.')

    doc.add_paragraph('\n\n_________________\nSignature of Donor').alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_heading('ENDORSEMENT BY THE ART BANK', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('I/we have personally explained to {full_name}, the details and implications of his signing this consent/approval form, and made sure to the extent humanly possible that he understands these details and implications.')

    # Left-align signature and address blocks
    for text in [
        '\n_______________________________\nName and signature of the Doctor',
        '\n___________________________________________________________\nName, address and signature of the Witness from the ART bank',
        '\nName and address of the ART bank\nCryoconserve Private Limited,\n3rd Floor, 59/1, Dr Rajkumar Road, 2nd Block,\nRajajinagar, Bengaluru, Karnataka 560010',
        '\nDated: {date}'
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save('templates_docx/form_15.docx')

def create_medical_history_template():
    """Generates the Medical History and Screening Report for Semen Donor."""
    doc = Document()
    apply_official_style(doc)

    heading = doc.add_heading('MEDICAL HISTORY AND SCREENING REPORT FOR SEMEN DONOR', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Left-align header info
    for label, value in [("ART Bank Name: ", "Cryoconserve"), ("Donor ID: ", "{donor_id}"), ("Date: ", "{date}")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(label).bold = True
        p.add_run(value)

    doc.add_heading('Section A: Donor Identification and Registration', level=2)
    doc.add_paragraph('1. Full Name: {full_name}')
    doc.add_paragraph('2. Date of Birth: {date_of_birth} (As per Aadhaar, Enclosed)')
    doc.add_paragraph('3. Contact Information:')
    doc.add_paragraph('   · Address: {address} (As per Aadhaar, Enclosed)', style='List Bullet')
    doc.add_paragraph('   · Phone Number: {contact_number}', style='List Bullet')
    doc.add_paragraph('   · Email: {email_address}', style='List Bullet')
    doc.add_paragraph('   · Aadhaar Number: {aadhaar_number}', style='List Bullet')

    doc.add_heading('Section B: Medical and Genetic Screening', level=2)
    doc.add_paragraph('1. Date of last comprehensive medical examination: {last_medical_exam}')
    doc.add_paragraph('2. Results of recent blood tests:')
    doc.add_paragraph('   · Human immunodeficiency virus (HIV), types 1 and 2: {hiv_results}', style='List Bullet 2')
    doc.add_paragraph('   · Hepatitis B virus (HBV): {hbv_results}', style='List Bullet 2')
    doc.add_paragraph('   · Hepatitis C virus (HCV): {hcv_results}', style='List Bullet 2')
    doc.add_paragraph('   · Treponema pallidum (syphilis) through VDRL: {vdrl_results}', style='List Bullet 2')
    doc.add_paragraph('3. Detailed family medical history, including any genetic conditions:\n{family_history}')
    doc.add_paragraph('4. Record of any serious illnesses or surgeries:\n{serious_illness}')
    doc.add_paragraph('5. Current medications and known allergies:\n{current_medications}, {allergies}')

    doc.add_heading('Section C: Consent for Cryopreservation and Use', level=2)
    doc.add_paragraph('1. Consent for cryopreservation of sperm: {consent_cryopreservation}')
    doc.add_paragraph('2. Consent for the use of sperm by ART Bank: {consent_art_bank}')

    doc.add_heading('Section D: National Registry Update Consent', level=2)
    doc.add_paragraph('1. Consent to update donor information in the National Registry: {consent_registry}')
    
    doc.add_paragraph('Declaration and Consent').runs[0].bold = True
    doc.add_paragraph('I hereby declare that the information provided above is true and complete to the best of my knowledge and I consent to the screening, collection, registration, and cryopreservation of my sperm as per the ART Regulation Act, 2021. I also consent to the maintenance of my records and the regular update of the National Registry as required by the Act. Furthermore, I declare I have never donated my semen to any ART clinic or bank, nor through any other means, and I will not donate my semen in the future.')
    
    doc.add_paragraph('\n\nSignature: _______________________________ Date: {date}').alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save('templates_docx/medical_history.docx')

def create_donor_info_template():
    """Generates the Sperm/Semen Donor Information Form."""
    doc = Document()
    apply_official_style(doc)

    doc.add_heading('SPERM/SEMEN DONOR INFORMATION FORM', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('Personal Details:', level=2)
    doc.add_paragraph('1. Full Name: {full_name}')
    doc.add_paragraph('2. Date of Birth: {date_of_birth} (As per Aadhaar, Attached)')
    doc.add_paragraph('3. Contact Number: {contact_number}')
    doc.add_paragraph('4. Email Address: {email_address}')
    doc.add_paragraph('5. Aadhaar Number: {aadhaar_number}')

    doc.add_heading('Health and Medical History:', level=2)
    doc.add_paragraph('1. Any Known Genetic Disorders or Medical Conditions: {genetic_disorders}')
    doc.add_paragraph('2. Family History of Genetic Conditions: {family_history}')
    doc.add_paragraph('3. Current Medications: {current_medications}')
    doc.add_paragraph('4. Allergies: {allergies}')

    doc.add_heading('Lifestyle and Habits:', level=2)
    doc.add_paragraph('1. Smoking Habits: {smoking}')
    doc.add_paragraph('   · Frequency: {smoking_frequency} (e.g., daily, occasionally)', style='List Bullet')
    doc.add_paragraph('   · Number of Cigarettes per Day: {cigarettes_per_day}', style='List Bullet')
    doc.add_paragraph('2. Alcohol Consumption: {alcohol}')
    doc.add_paragraph('   · Frequency: {alcohol_frequency} (e.g., weekly, monthly)', style='List Bullet')
    doc.add_paragraph('   · Average Amount Consumed: {alcohol_amount} (e.g., number of drinks)', style='List Bullet')
    doc.add_paragraph('3. Recreational Drug Use: {drug_use}')
    doc.add_paragraph('4. Dietary Preferences: {diet}')

    doc.add_heading('Reproductive History:', level=2)
    doc.add_paragraph('1. Marital Status: {marital_status}')
    doc.add_paragraph('2. Number of Biological Children (if any): {num_children}')
    doc.add_paragraph('{children_ages}')
    doc.add_paragraph('3. Previous Donor Experience (if applicable): {donor_experience}')
    doc.add_paragraph('   · Frequency of Donations (if known): {donation_frequency}', style='List Bullet')

    doc.add_heading('Physical Attributes:', level=2)
    doc.add_paragraph('1. Height: {height} (in cm/ft-in)')
    doc.add_paragraph('2. Weight: {weight} (in kg/lbs)')
    doc.add_paragraph('3. Educational Qualifications: {education}')
    doc.add_paragraph('4. Mother Tongue: {mother_tongue}')
    doc.add_paragraph('5. Skin Colour: {skin_colour}')
    doc.add_paragraph('6. Hair Colour: {hair_colour}')
    doc.add_paragraph('7. Eye Colour: {eye_colour}')
    doc.add_paragraph('8. Religion: {religion}')
    doc.add_paragraph('9. Occupation: {occupation}')

    doc.add_heading('Consent and Legal Acknowledgment:', level=2)
    doc.add_paragraph('1. I understand that my genetic material will be used for assisted reproductive purposes.')
    doc.add_paragraph('2. I consent to the storage and use of my sperm/semen for fertility treatments.')
    doc.add_paragraph('3. I acknowledge that I am voluntarily providing this information for legal purposes.')
    doc.add_paragraph('4. Furthermore, I declare I have never donated my semen to any ART clinic or bank, nor through any other means, and I will not donate my semen in the future.')
    
    doc.add_paragraph('\n\nSignature: ___________________________ Date: {date}').alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save('templates_docx/donor_info.docx')

def create_form_13_template():
    """Generates the Oocyte Donor Consent Form (Form 13)."""
    doc = Document()
    apply_official_style(doc)

    doc.add_heading('FORM 13', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('[Refer Rule 13 (f) (viii)]').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('Consent Form for the Donor of Oocytes', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run('I, ').bold = True
    p.add_run('{full_name}, residing at {address}, District: {district}, State: {state}, PIN Code: {pin_code}, Mobile: {contact_number}, Aadhaar Number {aadhaar_number}, willingly consent to donate my oocyte to couples who are unable to have a child by other means. At this stage and to the best of my knowledge I am free of any infectious diseases or genetic disorders.')

    p = doc.add_paragraph('I have had a full discussion with ')
    p.add_run('{doctor_name}').bold = True
    p.add_run(' on {date_of_discussion}, at {ivf_name}, {ivf_address}.')
    
    doc.add_paragraph('(I understand that there will be no direct or indirect contact between me and the recipient, and my personal identity will not be disclosed to the recipient or to the child born through the use of my gamete: If applicable)')
    doc.add_paragraph('I understand that I shall have no rights whatsoever on the resulting offspring and vice versa.')
    doc.add_paragraph('I understand that the method of treatment may include:')
    doc.add_paragraph('Stimulating my ovaries for multifollicular development.', style='List Number')
    doc.add_paragraph('The recovery of one or more of my eggs under ultrasound-guidance or by laparoscopy under sedation or general anesthesia.', style='List Number')
    doc.add_paragraph('The fertilization of my oocytes with recipient’s husband’s or donor sperm and transferring the resulting embryo into the recipient.', style='List Number')
    doc.add_paragraph('I understand and accept that the drugs that are used to stimulate the ovaries have temporary side-effects like nausea, headaches, and abdominal bloating. In a small proportion of cases, a condition called ovarian hyperstimulation occurs. Such cases can be identified in time but only to a limited extent. Further, at times the ovarian response is poor or absent in spite of using a high dose of drugs. Under these circumstances, the treatment cycle will be cancelled.')

    p = doc.add_paragraph('\n')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run('Name, address and signature of woman\n').bold = True
    p.add_run('{full_name}\nAddress: {address}, District: {district}, State: {state}, PIN Code: {pin_code}')

    doc.add_heading('Endorsement by the ART Clinic', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('We have personally explained to {full_name}, the details and implications of her signing this consent/approval form, and made sure to the extent humanly possible that she understands these details and implications.')

    # Left-align signature and address blocks
    for text in [
        '\nName, address and signature of the Witness from the clinic\nMrs. Ruby Stella\nIVF Access, Coimbatore, at 609, 2nd Floor, Avinashi Road, above Pazhamudir Plus, Peelamedu, Coimbatore, Tamil Nadu 641004',
        '\nName and signature of the Doctor\n{doctor_name}',
        '\nName and address of the ART clinic\n{ivf_name}, {ivf_address}',
        '\nName and address of the ART bank that recruited and screened the donor\nCryoconserve, 3rd Floor, 59/1, 2nd Block, Rajajinagar, Bengaluru-560010',
        '\nDated: {date}'
    ]:
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save('templates_docx/form_13.docx')
    
def create_oocyte_medical_history_template():
    """Generates the Medical History and Screening Report for Oocyte Donor."""
    doc = Document()
    apply_official_style(doc)

    doc.add_heading('MEDICAL HISTORY AND SCREENING REPORT FOR OOCYTE DONOR', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for label, value in [("ART Bank Name: ", "Cryoconserve"), ("Donor ID: ", "{donor_id}"), ("Date: ", "{date}")]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(label).bold = True
        p.add_run(value)

    doc.add_heading('Section A: Donor Identification and Registration', level=2)
    doc.add_paragraph('1. Full Name: {full_name}')
    doc.add_paragraph('2. Date of Birth: {date_of_birth} (As per Aadhaar, Enclosed)')
    doc.add_paragraph('3. Marital Status: {marital_status}')
    doc.add_paragraph('4. Contact Information:')
    doc.add_paragraph('   · Address: {address}, District: {district}, State: {state}, PIN Code: {pin_code}', style='List Bullet')
    doc.add_paragraph('   · Phone Number: {contact_number}', style='List Bullet')
    doc.add_paragraph('   · Email: {email_address}', style='List Bullet')
    doc.add_paragraph('   · Aadhaar Number: {aadhaar_number}', style='List Bullet')

    doc.add_heading('Section B: Medical and Genetic Screening', level=2)
    doc.add_paragraph('1. Date of last comprehensive medical examination: {last_medical_exam}')
    doc.add_paragraph('2. Results of recent blood tests:')
    doc.add_paragraph('   · Human immunodeficiency virus (HIV), types 1 and 2: {hiv_results}', style='List Bullet 2')
    doc.add_paragraph('   · Hepatitis B virus (HBV): {hbv_results}', style='List Bullet 2')
    doc.add_paragraph('   · Hepatitis C virus (HCV): {hcv_results}', style='List Bullet 2')
    doc.add_paragraph('   · Treponema pallidum (syphilis) through VDRL: {vdrl_results}', style='List Bullet 2')
    doc.add_paragraph('3. Detailed family medical history, including any genetic conditions:\n{family_history}')
    doc.add_paragraph('4. Record of any serious illnesses or surgeries:\n{serious_illness}')
    doc.add_paragraph('5. Current medications and known allergies:\n{current_medications}, {allergies}')

    doc.add_heading('Section C: Ovarian Reserve Assessment', level=2)
    doc.add_paragraph('1. Antral follicle count (AFC): {antral_follicle_count}')
    doc.add_paragraph('2. Follicle-stimulating hormone (FSH) levels: {fsh_levels}')
    doc.add_paragraph('3. Anti-Müllerian hormone (AMH) levels: {amh_levels}')

    doc.add_heading('Section D: Lifestyle and Health Assessment', level=2)
    doc.add_paragraph('1. Tobacco Use: {tobacco_use}')
    doc.add_paragraph('   · If yes, Frequency: {tobacco_frequency}', style='List Bullet')
    doc.add_paragraph('2. Alcohol Consumption: {alcohol}')
    doc.add_paragraph('   · If yes, Frequency: {alcohol_frequency}', style='List Bullet')
    doc.add_paragraph('3. Recreational Drug Use: {drug_use}')
    doc.add_paragraph('4. Exercise Routine: {exercise_routine}')

    doc.add_heading('Section E: National Registry Update Consent', level=2)
    doc.add_paragraph('1. Consent to update donor information in the National Registry: {consent_registry}')
    
    doc.add_paragraph('Declaration and Consent').runs[0].bold = True
    doc.add_paragraph('I hereby declare that the information provided above is true and complete to the best of my knowledge. I consent to the screening, collection, registration, and cryopreservation of my oocytes as per the ART Regulation Act, 2021. I also consent to the maintenance of my records and the regular update of the National Registry as required by the Act.')
    
    doc.add_paragraph('\n\nSignature: _______________________________ Date: {date}').alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.save('templates_docx/oocyte_medical_history.docx')

def create_oocyte_donor_affidavit_template():
    """Generates the Oocyte Donor Affidavit/Declaration."""
    doc = Document()
    apply_official_style(doc)
    
    doc.add_heading('OOCYTE DONOR AFFIDAVIT / DECLARATION', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('I, ').bold = True
    p.add_run('{full_name}, aged {age} Years (DOB {date_of_birth}), residing at {address}, District: {district}, State: {state}, PIN Code: {pin_code}, Mobile: {contact_number}, Aadhaar Number: {aadhaar_number}, do hereby state on solemn affirmation as under:')
    
    doc.add_paragraph('At this stage, I have {num_children} Child(ren) more than 2-year-old. I have presented supportive documents of my children. I also state that I will not donate my oocyte to any other couple. I also declare that I have not undergone any oocyte donation before.')
    
    p = doc.add_paragraph('I have been counselled by ')
    p.add_run('{doctor_name}').bold = True
    p.add_run(' at {ivf_name}, {ivf_address}. I have been explained about the process of oocyte retrieval procedure.')
    
    doc.add_paragraph('I, willingly without any financial interest consent to donate my oocyte to couple/individual who are unable to have a child by other means, and we are not taking any kind of donation/amount for this.')
    doc.add_paragraph('I, declare that above given information is true to my knowledge and all the responsibility of any legal liabilities will be upon me.')

    doc.add_paragraph('\nPlace: {place}').alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph('Date: {date}').alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph('\n\nSignature of Donor').alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save('templates_docx/oocyte_donor_affidavit.docx')

def create_commissioning_couple_affidavit_template():
    """Generates the Commissioning Couple's Affidavit."""
    doc = Document()
    apply_official_style(doc)

    doc.add_heading('COMMISSIONING COUPLE\'S AFFIDAVIT', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('We, ').bold = True
    p.add_run('{female_name}, W/o {male_name}, aged about {female_age} Years (DOB {female_dob}), Aadhaar Number {female_aadhaar}, Occupation {female_occupation}, and {male_name}, aged about {male_age} Years (DOB {male_dob}), both residing at {address}, District: {district}, State: {state}, PIN Code: {pin_code}, do hereby solemnly affirm and depose as under:')
    
    text_paras = [
        'We state that, we are unable to conceive a child and have decided to opt for infertility treatment and thereby approached {ivf_name}, {ivf_address} (hereinafter referred to as "ART Clinic"). However, we agreed to avail the ART treatment using the Oocytes from the Oocyte Donor, we the intending couple shall obtain an insurance coverage in favor of Oocyte Donor and give a guarantee by signing an affidavit as prescribed under the Assisted Reproductive Technology (Regulation) Act 2021 and Assisted Reproductive Technology (Regulation) Rules 2022.',
        'That, in view of the above-mentioned ART Act and Rules, we hereby sign this Affidavit in compliance with Rule 12 (ii) of the Assisted Reproductive Technology (Regulation) Rules 2022 (in short "ART Rules").',
        'Further, we hereby undertake to bear the full medical expenditures for the Oocyte Donor during the treatment and provide guarantee of compensation for specified loss, damage, complication or death of oocyte donor during the process of oocyte retrieval as prescribed under Section 22(1)(b) and Section 22(4)(ii) of Assisted Reproductive Technology (Regulation) Act, 2021 and 12(ii) of the Assisted Reproductive Technology Act (Regulation) 2021.',
        'We abide to take a general health insurance coverage and a life insurance coverage in favor of oocyte donor for the period of Twelve (12) Months from an Insurance Company (recognized by Insurance Regulatory Development Authority) for an amount sufficient to cover all expenses for all complications arising due to oocyte retrieval as per 22(1)(b) of the Assisted Reproductive Technology (Regulation) Act 2021, and 12(i) of Assisted Reproductive Technology (Regulation) Rules, 2022.',
        'We also state that this Affidavit is signed by both of us with free consent and without any undue influence and coercion.',
        'That the contents stated in para 1 to para 5 of this Affidavit are true and correct to the best of our knowledge and also, we indemnify and keep the ART Bank and the ART Clinic Indemnified, if any statement or fact stated by us is incorrect or false.'
    ]
    for text in text_paras:
        doc.add_paragraph(text, style='List Number')

    doc.add_paragraph('\n\nDEPONENT').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading('VERIFICATION', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Verified at {place} on {date} that the contents of the affidavit in Para-1 to 6 are true and correct to the best of my knowledge and nothing has been concealed therefrom.')
    
    doc.add_paragraph('\nPlace: {place}').alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph('Date: {date}').alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph('\n\nDEPONENT').alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save('templates_docx/commissioning_couple_affidavit.docx')

def create_all_templates():
    """Creates the output directory and generates all DOCX templates."""
    output_dir = 'templates_docx'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        print(f"Created directory: {output_dir}")
    
    templates_to_create = {
        "Form 15": create_form_15_template,
        "Semen Donor Medical History": create_medical_history_template,
        "Semen Donor Info Form": create_donor_info_template,
        "Form 13": create_form_13_template,
        "Oocyte Donor Medical History": create_oocyte_medical_history_template,
        "Oocyte Donor Affidavit": create_oocyte_donor_affidavit_template,
        "Commissioning Couple Affidavit": create_commissioning_couple_affidavit_template,
    }

    print("Starting template generation...")
    for name, func in templates_to_create.items():
        try:
            func()
            print(f"  ✓ Successfully generated: {name}")
        except Exception as e:
            print(f"  ✗ Failed to generate {name}: {e}")
    
    print("\nAll templates have been processed.")

if __name__ == '__main__':
    create_all_templates()